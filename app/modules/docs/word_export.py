from flask import current_app
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from datetime import datetime
import os
import re
from bs4 import BeautifulSoup
from io import BytesIO

def export_document_to_word(document, organization=None):
    """Export a document to Word format (supports both markdown and HTML, with images)"""
    # Create a new Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Convert content to HTML based on content_type
    if document.content_type == 'html':
        html_content = document.content or ''
    else:
        # Convert markdown to HTML
        html_content = markdown.markdown(document.content or '', extensions=['extra', 'codehilite'])
    
    # Add header with document title
    title_para = doc.add_heading(document.title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add metadata
    if organization:
        meta_para = doc.add_paragraph()
        meta_para.add_run('Organization: ').bold = True
        meta_para.add_run(organization.name)
    
    created_para = doc.add_paragraph()
    created_para.add_run('Created: ').bold = True
    created_para.add_run(document.created_at.strftime('%Y-%m-%d %H:%M:%S') if document.created_at else 'N/A')
    
    updated_para = doc.add_paragraph()
    updated_para.add_run('Last Updated: ').bold = True
    updated_para.add_run(document.updated_at.strftime('%Y-%m-%d %H:%M:%S') if document.updated_at else 'N/A')
    
    # Add a separator
    doc.add_paragraph()
    
    # Parse HTML and convert to Word document
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Find body tag if it exists, otherwise use the root
    body = soup.find('body')
    root = body if body else soup
    
    # Process each element
    for element in root.children:
        if hasattr(element, 'name') and element.name:
            process_element(doc, element, current_app)
        elif isinstance(element, str) and element.strip():
            # Handle standalone text nodes
            para = doc.add_paragraph()
            para.add_run(element.strip())
    
    # Add footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(f'Exported on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(127, 127, 127)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def process_element(doc, element, app):
    """Process an HTML element and add it to the Word document"""
    tag_name = element.name.lower() if element.name else None
    
    if tag_name == 'h1':
        para = doc.add_heading(element.get_text(), level=1)
    elif tag_name == 'h2':
        para = doc.add_heading(element.get_text(), level=2)
    elif tag_name == 'h3':
        para = doc.add_heading(element.get_text(), level=3)
    elif tag_name == 'h4':
        para = doc.add_heading(element.get_text(), level=4)
    elif tag_name == 'h5':
        para = doc.add_heading(element.get_text(), level=5)
    elif tag_name == 'h6':
        para = doc.add_heading(element.get_text(), level=6)
    elif tag_name == 'p':
        para = doc.add_paragraph()
        process_inline_elements(para, element, app)
    elif tag_name == 'ul':
        process_list(doc, element, app, False)
    elif tag_name == 'ol':
        process_list(doc, element, app, True)
    elif tag_name == 'img':
        process_image(doc, element, app)
    elif tag_name == 'table':
        process_table(doc, element, app)
    elif tag_name == 'blockquote':
        para = doc.add_paragraph()
        para.style = 'Quote'
        process_inline_elements(para, element, app)
    elif tag_name == 'pre':
        para = doc.add_paragraph()
        para.style = 'No Spacing'
        run = para.add_run(element.get_text())
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    elif tag_name == 'code':
        # Inline code - handled in process_inline_elements
        pass
    elif tag_name in ['div', 'span', 'section', 'article']:
        # Process children recursively
        # If it's a span, try to process inline first
        if tag_name == 'span':
            # For span, try to create a paragraph if there's content
            has_block_children = any(hasattr(c, 'name') and c.name and c.name.lower() in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'table'] for c in element.children)
            if not has_block_children:
                # Treat as inline content
                para = doc.add_paragraph()
                process_inline_elements(para, element, app)
            else:
                # Has block children, process them
                for child in element.children:
                    if hasattr(child, 'name') and child.name:
                        process_element(doc, child, app)
                    elif isinstance(child, str) and child.strip():
                        para = doc.add_paragraph()
                        para.add_run(child.strip())
        else:
            # For div/section/article, process children
            for child in element.children:
                if hasattr(child, 'name') and child.name:
                    process_element(doc, child, app)
                elif isinstance(child, str) and child.strip():
                    para = doc.add_paragraph()
                    para.add_run(child.strip())
    elif tag_name == 'br':
        doc.add_paragraph()
    elif tag_name in ['script', 'style']:
        # Skip script and style tags
        pass
    else:
        # For unknown tags, try to process as paragraph
        if element.get_text(strip=True):
            para = doc.add_paragraph()
            process_inline_elements(para, element, app)

def process_inline_elements(para, element, app):
    """Process inline elements (text formatting, links, images, etc.)"""
    from bs4 import NavigableString
    
    for content in element.children:
        if isinstance(content, (str, NavigableString)):
            # Plain text
            text = str(content).strip()
            if text:
                para.add_run(text)
        elif hasattr(content, 'name') and content.name:
            tag_name = content.name.lower()
            if tag_name == 'strong' or tag_name == 'b':
                run = para.add_run(content.get_text())
                run.bold = True
            elif tag_name == 'em' or tag_name == 'i':
                run = para.add_run(content.get_text())
                run.italic = True
            elif tag_name == 'u':
                run = para.add_run(content.get_text())
                run.underline = True
            elif tag_name == 'code':
                run = para.add_run(content.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif tag_name == 'a':
                href = content.get('href', '')
                link_text = content.get_text()
                if href:
                    add_hyperlink(para, href, link_text)
                else:
                    para.add_run(link_text)
            elif tag_name == 'img':
                process_image_inline(para, content, app)
            elif tag_name == 'br':
                para.add_run('\n')
            else:
                # Recursively process nested elements
                process_inline_elements(para, content, app)

def process_list(doc, element, app, ordered=False):
    """Process a list (ul or ol)"""
    for li in element.find_all('li', recursive=False):
        para = doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')
        process_inline_elements(para, li, app)

def process_image(doc, element, app):
    """Process an image element as a block element"""
    src = element.get('src', '')
    if not src:
        return
    
    # Get absolute file path
    image_path = get_image_path(src, app)
    if not image_path or not os.path.exists(image_path):
        # If image not found, add a placeholder text
        para = doc.add_paragraph()
        para.add_run(f'[Image: {src}]').italic = True
        return
    
    try:
        # Add image to document
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        # Add image with max width of 6 inches
        run.add_picture(image_path, width=Inches(6))
    except Exception as e:
        # If image can't be added, add placeholder
        para = doc.add_paragraph()
        para.add_run(f'[Image: {src} - Error: {str(e)}]').italic = True

def process_image_inline(para, element, app):
    """Process an image element inline"""
    src = element.get('src', '')
    if not src:
        return
    
    # Get absolute file path
    image_path = get_image_path(src, app)
    if not image_path or not os.path.exists(image_path):
        # If image not found, add a placeholder text
        para.add_run(f'[Image: {src}]').italic = True
        return
    
    try:
        # Add image inline
        run = para.add_run()
        # Add image with max width of 4 inches for inline
        run.add_picture(image_path, width=Inches(4))
    except Exception as e:
        # If image can't be added, add placeholder
        para.add_run(f'[Image: {src} - Error: {str(e)}]').italic = True

def process_table(doc, element, app):
    """Process a table element"""
    rows = element.find_all('tr')
    if not rows:
        return
    
    # Determine number of columns
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols)
    
    if max_cols == 0:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            if col_idx < max_cols:
                cell_para = table.rows[row_idx].cells[col_idx].paragraphs[0]
                cell_para.clear()
                # Check if it's a header cell
                if cell.name == 'th':
                    run = cell_para.add_run(cell.get_text())
                    run.bold = True
                else:
                    process_inline_elements(cell_para, cell, app)

def get_image_path(src, app):
    """Convert image src to absolute file path"""
    # If it's a data URL, return None (can't embed data URLs in Word easily)
    if src.startswith('data:'):
        return None
    
    # If it's a relative URL starting with /static/
    if src.startswith('/static/'):
        static_folder = app.static_folder
        relative_path = src.replace('/static/', '')
        absolute_path = os.path.abspath(os.path.join(static_folder, relative_path))
        return absolute_path if os.path.exists(absolute_path) else None
    
    # If it's already an absolute path
    if os.path.isabs(src) and os.path.exists(src):
        return src
    
    # If it's a file:// URL
    if src.startswith('file://'):
        path = src.replace('file://', '').replace('file:///', '')
        # Handle Windows paths
        if os.name == 'nt' and len(path) > 1 and path[1] == ':':
            return path
        elif os.path.exists(path):
            return path
    
    return None

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    # This creates the hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create the run element for the hyperlink text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink style (blue, underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

